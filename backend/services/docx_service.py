from io import BytesIO
from typing import Optional, Tuple

from fastapi import HTTPException
import docx

from backend.logger import logger


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from a DOCX file in memory."""
    try:
        doc = docx.Document(BytesIO(file_bytes))
        text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
        return text
    except Exception as e:
        logger.error("DOCX extraction failed: %s", e)
        raise HTTPException(400, f"DOCX processing error: {str(e)}")


def extract_text_from_docx_with_stats(file_bytes: bytes) -> Tuple[str, Optional[int]]:
    """Extract text from a DOCX and return (text, pages_seen).

    DOCX has no reliable page count without rendering; returns None.
    """
    try:
        doc = docx.Document(BytesIO(file_bytes))
        text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
        return text, None
    except Exception as e:
        logger.error("DOCX extraction failed: %s", e)
        raise HTTPException(400, f"DOCX processing error: {str(e)}")
